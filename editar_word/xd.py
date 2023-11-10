from docx import Document
from docx.shared import Pt
import os
# import module
from pdf2image import convert_from_path


# 1. Abre el archivo existente
doc = Document('/editar_word/plantilla_boleta.docx')
nombre="A"
# 2. Realiza las modificaciones que desees
for paragraph in doc.paragraphs:
    # Ejemplo: Modificar el contenido del primer párrafo
    if "#corre" in paragraph.text:
        paragraph.text = paragraph.text.replace("#corre","Prpgramacion")
        # paragraph.bold = False
    
    if "#control" in paragraph.text:
        paragraph.text = paragraph.text.replace("#control", "21301061550024")
    
    if "#nombre" in paragraph.text:
        paragraph.text = paragraph.text.replace("#nombre", nombre)
    if "#semestre" in paragraph.text:
        paragraph.text = paragraph.text.replace("#semestre", "5")
    
    if "#turno" in paragraph.text:
        paragraph.text = paragraph.text.replace("#turno", "Matutino")
   


    palabra_a_buscar="NOMBRE"
    
    for paragraph in doc.paragraphs:
        if palabra_a_buscar in paragraph.text:
            for run in paragraph.runs:
                if palabra_a_buscar in run.text:
                    # Encuentra la palabra a buscar y aplica formato en negritas
                    run.text = run.text.replace(palabra_a_buscar, f"{palabra_a_buscar}", 1)
                    run.bold = True
    
                    


#rellenado de tabla , calificaciones
tabla = doc.tables[0] 
for j in range(1,8):
    fila = tabla.rows[j]
    for i in range(0,8):
        fila.cells[i].text=str(j),str(i)
        
    
 # Guarda el documento modificado
doc.save(nombre+'.docx')


#transformar el documento de docx a pdf bien clean xd papu pro
from docx2pdf import convert

inputFile = nombre+'.docx'
outputFile = nombre+'.pdf'

convert(inputFile, outputFile)


#eliminar el documento de docx despues de haberlo escrito en pdf
import os
os.remove(nombre+'.docx')
